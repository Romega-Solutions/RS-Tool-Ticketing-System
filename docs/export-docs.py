"""
Export plan docs to DOCX for offline sharing.
Usage: python export-docs.py
Output: docs/exports/*.docx
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PLAN_DIR = Path(__file__).parent / "plan"
EXPORT_DIR = Path(__file__).parent / "exports"

# RS branding
RS_BLUE = RGBColor(0x00, 0x69, 0xD9)
RS_DARK = RGBColor(0x33, 0x33, 0x33)

# Docs to export (in order)
DOCS = [
    "executive-summary.md",
    "product-vision.md",
    "feasibility.md",
    "features.md",
    "data-model.md",
    "screens.md",
    "weekly-report-workflow.md",
    "migration.md",
    "plane-configuration.md",
    "presentation-content.md",
]


def style_heading(paragraph, level=1):
    for run in paragraph.runs:
        run.font.color.rgb = RS_BLUE
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)


def add_table_from_rows(doc, header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Table Grid"

    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    return table


def parse_table_line(line):
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def is_separator_line(line):
    return bool(re.match(r"^\s*\|[\s\-:|]+\|\s*$", line))


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RS_DARK

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            style_heading(p, 1)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            style_heading(p, 2)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            style_heading(p, 3)
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Table
        if "|" in line and not is_separator_line(line):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                if not is_separator_line(lines[i]):
                    table_lines.append(parse_table_line(lines[i]))
                i += 1
            if len(table_lines) >= 2:
                add_table_from_rows(doc, table_lines[0], table_lines[1:])
                doc.add_paragraph("")  # spacing
            elif len(table_lines) == 1:
                add_table_from_rows(doc, table_lines[0], [])
                doc.add_paragraph("")
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        # Checkbox list items
        if re.match(r"^-\s*\[([ xX])\]\s", line):
            checked = line[3] in ("x", "X")
            text = line[5:].strip()
            marker = "☑" if checked else "☐"
            p = doc.add_paragraph(f"{marker} {text}")
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            text = m.group(2).strip()
            p = doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Indented sub-items
        if line.startswith("  - ") or line.startswith("  * "):
            text = line.strip().lstrip("- ").lstrip("* ")
            p = doc.add_paragraph(f"  • {text}")
            p.paragraph_format.left_indent = Inches(0.6)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph — handle inline bold/italic
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1

    doc.save(str(docx_path))


def add_formatted_text(paragraph, text):
    # Split on **bold** and *italic* patterns
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def main():
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    exported = []
    for md_name in DOCS:
        md_path = PLAN_DIR / md_name
        if not md_path.exists():
            print(f"  SKIP: {md_name} (not found)")
            continue

        docx_name = md_name.replace(".md", ".docx")
        docx_path = EXPORT_DIR / docx_name
        convert_md_to_docx(md_path, docx_path)
        exported.append(docx_name)
        print(f"  OK: {docx_name}")

    print(f"\nExported {len(exported)} docs to: {EXPORT_DIR}")


if __name__ == "__main__":
    main()
