"""
Export key RS Ticketing System docs to DOCX format.
Usage: python export-to-docx.py
Output: docs/exports/ folder with .docx files
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PLAN_DIR = Path(__file__).parent / "plan"
EXPORT_DIR = Path(__file__).parent / "exports"

# Key docs to export (filename, display title)
DOCS = [
    ("executive-summary.md", "RS Ticketing System — Executive Summary"),
    ("feasibility.md", "RS Ticketing System — Feasibility Research"),
    ("plane-configuration.md", "Plane.so Configuration Guide for Romega Solutions"),
    ("handoff-ken.md", "Ken's Deployment Handoff — RS Ticketing System"),
    ("product-vision.md", "RS Ticketing System — Product Vision"),
    ("features.md", "RS Ticketing System — Feature Breakdown"),
    ("weekly-report-workflow.md", "RS Ticketing System — Weekly Report Workflow"),
    ("migration.md", "RS Ticketing System — Migration Plan"),
]

RS_BLUE = RGBColor(0x00, 0x69, 0xD9)
RS_GRAY = RGBColor(0x33, 0x33, 0x33)


def style_document(doc):
    """Apply RS branding to document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RS_GRAY

    for i in range(1, 4):
        heading = doc.styles[f"Heading {i}"]
        heading.font.color.rgb = RS_BLUE
        heading.font.name = "Calibri"


def md_to_docx(md_path, title):
    """Convert a markdown file to a styled DOCX document."""
    doc = Document()
    style_document(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Code blocks
        if line.startswith("```"):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Normal"]
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Table rows
        if line.startswith("|") and not line.startswith("|--") and not re.match(r"^\|[\s\-:|]+\|$", line):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            i += 1
            # Check if next line is separator or another row
            if i < len(lines):
                next_line = lines[i].rstrip("\n")
                if not next_line.startswith("|") or re.match(r"^\|[\s\-:|]+\|$", next_line):
                    if re.match(r"^\|[\s\-:|]+\|$", next_line):
                        i += 1  # skip separator
                    # Check for more data rows
                    while i < len(lines) and lines[i].rstrip("\n").startswith("|"):
                        row_line = lines[i].rstrip("\n")
                        if re.match(r"^\|[\s\-:|]+\|$", row_line):
                            i += 1
                            continue
                        cells = [c.strip() for c in row_line.split("|")[1:-1]]
                        if cells:
                            table_rows.append(cells)
                        i += 1

                    # Render table
                    if table_rows:
                        max_cols = max(len(r) for r in table_rows)
                        tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                        tbl.style = "Light Grid Accent 1"
                        for ri, row in enumerate(table_rows):
                            for ci, cell in enumerate(row):
                                if ci < max_cols:
                                    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
                                    clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)
                                    clean = re.sub(r"`(.+?)`", r"\1", clean)
                                    tbl.rows[ri].cells[ci].text = clean
                        table_rows = []
                    continue
            # If we still have rows pending, flush them
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.style = "Light Grid Accent 1"
                for ri, row in enumerate(table_rows):
                    for ci, cell in enumerate(row):
                        if ci < max_cols:
                            tbl.rows[ri].cells[ci].text = cell
                table_rows = []
            continue

        # Skip table separator lines
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Horizontal rules
        if line.startswith("---"):
            i += 1
            continue

        # Bullet points
        if re.match(r"^[-*] ", line):
            text = re.sub(r"^[-*] ", "", line)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered lists
        if re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            p = doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraphs
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        if text.strip():
            doc.add_paragraph(text)
        i += 1

    return doc


def main():
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    exported = []
    for filename, title in DOCS:
        md_path = PLAN_DIR / filename
        if not md_path.exists():
            print(f"  SKIP: {filename} (not found)")
            continue

        doc = md_to_docx(md_path, title)
        out_name = filename.replace(".md", ".docx")
        out_path = EXPORT_DIR / out_name
        doc.save(str(out_path))
        exported.append(out_name)
        print(f"  OK: {out_name}")

    print(f"\nExported {len(exported)} docs to {EXPORT_DIR}/")


if __name__ == "__main__":
    main()
